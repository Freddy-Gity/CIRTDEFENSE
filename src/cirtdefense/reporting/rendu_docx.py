"""Rendu Word : la version reprise et annotée.

Le PDF est destiné à la signature ; le Word est destiné au travail. Un chef
de service y ajoute un paragraphe, un juriste y insère une réserve, une
assistante y colle le rapport dans un courrier plus large. Le fichier doit
donc rester structuré — de vrais styles de titre, de vrais tableaux — et non
livrer une image du document.

python-docx est retenu pour la même raison que ReportLab : bibliothèque en
Python pur, sans dépendance système, utilisable sur un site coupé du réseau.
"""

from __future__ import annotations

import io
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .document import (
    Document,
    Encadre,
    Graphique,
    Liste,
    Paragraphe,
    SautDePage,
    Tableau,
    Titre,
)

VERT_ADMIN = RGBColor(0x1B, 0x5E, 0x3A)
GRIS_TEXTE = RGBColor(0x55, 0x5B, 0x5E)
ROUGE_ALERTE = RGBColor(0xA0, 0x20, 0x20)
ORANGE_ATTENTION = RGBColor(0xA0, 0x6A, 0x10)

LOGO_PAR_DEFAUT = Path("web/static/logo-antic.png")

LARGEUR_BARRE = 26
"""Longueur de la piste du diagramme, en caractères pleins."""


def rendre(document: Document, logo: str | Path | None = None) -> bytes:
    fichier = DocxDocument()
    _mise_en_page(fichier)
    _titulature(fichier, document, _chemin_logo(logo, document))
    _cartouche(fichier, document)
    for bloc in document.blocs:
        _bloc(fichier, bloc)
    _signature(fichier, document)

    tampon = io.BytesIO()
    fichier.save(tampon)
    return tampon.getvalue()


# ------------------------------------------------------------ mise en page


def _mise_en_page(fichier) -> None:
    section = fichier.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = fichier.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)

    pied = section.footer.paragraphs[0]
    pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trace = pied.add_run("Centre de réponse aux incidents informatiques — ANTIC")
    trace.font.size = Pt(7.5)
    trace.font.color.rgb = GRIS_TEXTE


def _chemin_logo(logo: str | Path | None, document: Document) -> Path | None:
    for candidat in (logo, document.entete.logo, LOGO_PAR_DEFAUT):
        if candidat and Path(candidat).is_file():
            return Path(candidat)
    return None


def _titulature(fichier, document: Document, logo: Path | None) -> None:
    gauche, droite = document.entete.colonnes()
    tableau = fichier.add_table(rows=1, cols=3)
    tableau.alignment = WD_TABLE_ALIGNMENT.CENTER
    cellules = tableau.rows[0].cells
    tableau.columns[0].width = Cm(6.2)
    tableau.columns[1].width = Cm(4.6)
    tableau.columns[2].width = Cm(6.2)

    _colonne_titulature(cellules[0], gauche)
    _colonne_titulature(cellules[2], droite)

    milieu = cellules[1].paragraphs[0]
    milieu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo is not None:
        milieu.add_run().add_picture(str(logo), width=Cm(2.6))
    else:
        reserve = milieu.add_run("[ EMBLÈME ]")
        reserve.font.size = Pt(7)
        reserve.font.color.rgb = GRIS_TEXTE

    _filet(fichier.add_paragraph())


def _colonne_titulature(cellule, lignes: list[str]) -> None:
    cellule.paragraphs[0].text = ""
    premier = True
    for ligne in lignes:
        if ligne.startswith("*"):
            continue
        paragraphe = cellule.paragraphs[0] if premier else cellule.add_paragraph()
        premier = False
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraphe.paragraph_format.space_after = Pt(1)
        trace = paragraphe.add_run(ligne)
        trace.bold = True
        trace.font.size = Pt(7)


def _cartouche(fichier, document: Document) -> None:
    titre = fichier.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trace = titre.add_run(document.titre)
    trace.bold = True
    trace.font.size = Pt(15)
    trace.font.color.rgb = VERT_ADMIN

    cadre = fichier.add_table(rows=1, cols=1)
    cellule = cadre.rows[0].cells[0]
    _fond(cellule, "F0F2F3")
    cellule.paragraphs[0].text = ""
    for etiquette, valeur in (
        ("Référence", document.reference),
        ("Objet", document.objet),
        (
            "Établi le",
            f"{_date(document.etabli_le)} — par : {document.etabli_par}",
        ),
    ):
        paragraphe = (
            cellule.paragraphs[0]
            if not cellule.paragraphs[0].runs
            else cellule.add_paragraph()
        )
        paragraphe.paragraph_format.space_after = Pt(2)
        gras = paragraphe.add_run(f"{etiquette} : ")
        gras.bold = True
        gras.font.size = Pt(9.5)
        texte = paragraphe.add_run(str(valeur))
        texte.font.size = Pt(9.5)
    fichier.add_paragraph()


# ------------------------------------------------------------------ blocs


def _bloc(fichier, bloc: object) -> None:
    match bloc:
        case Titre():
            paragraphe = fichier.add_paragraph()
            paragraphe.paragraph_format.space_before = Pt(11)
            paragraphe.paragraph_format.space_after = Pt(4)
            trace = paragraphe.add_run(
                bloc.intitule.upper() if bloc.niveau <= 1 else bloc.intitule
            )
            trace.bold = True
            trace.font.size = Pt(12 if bloc.niveau <= 1 else 10.5)
            if bloc.niveau <= 1:
                trace.font.color.rgb = VERT_ADMIN
        case Paragraphe():
            paragraphe = fichier.add_paragraph()
            paragraphe.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            trace = paragraphe.add_run(bloc.texte)
            if bloc.accent:
                trace.bold = True
                trace.font.color.rgb = VERT_ADMIN
        case Liste():
            for index, element in enumerate(bloc.elements, 1):
                paragraphe = fichier.add_paragraph(
                    element, style="List Number" if bloc.numerotee else "List Bullet"
                )
                paragraphe.paragraph_format.space_after = Pt(2)
                del index
        case Tableau():
            _tableau(fichier, bloc)
        case Graphique():
            _graphique(fichier, bloc)
        case Encadre():
            _encadre(fichier, bloc)
        case SautDePage():
            fichier.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _tableau(fichier, bloc: Tableau) -> None:
    if not bloc.entetes:
        return
    tableau = fichier.add_table(rows=1, cols=len(bloc.entetes))
    tableau.style = "Table Grid"
    tableau.alignment = WD_TABLE_ALIGNMENT.CENTER

    for index, entete in enumerate(bloc.entetes):
        cellule = tableau.rows[0].cells[index]
        _fond(cellule, "1B5E3A")
        cellule.paragraphs[0].text = ""
        trace = cellule.paragraphs[0].add_run(entete)
        trace.bold = True
        trace.font.size = Pt(8.5)
        trace.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for numero, ligne in enumerate(bloc.lignes):
        cellules = tableau.add_row().cells
        for index, valeur in enumerate(ligne):
            if index >= len(cellules):
                break
            cellule = cellules[index]
            if numero % 2:
                _fond(cellule, "F0F2F3")
            cellule.paragraphs[0].text = ""
            if bloc.alignement(index) == "droite":
                cellule.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            trace = cellule.paragraphs[0].add_run(str(valeur))
            trace.font.size = Pt(8.5)

    if bloc.legende:
        legende = fichier.add_paragraph()
        legende.alignment = WD_ALIGN_PARAGRAPH.CENTER
        trace = legende.add_run(bloc.legende)
        trace.italic = True
        trace.font.size = Pt(8)
        trace.font.color.rgb = GRIS_TEXTE
    else:
        fichier.add_paragraph()


def _graphique(fichier, bloc: Graphique) -> None:
    """Le diagramme est composé en caractères pleins dans une police fixe.

    Insérer une image obligerait à embarquer un moteur de tracé et à figer le
    graphique ; en caractères, il reste modifiable dans Word comme le reste
    du document.
    """
    if not bloc.valeurs:
        return
    titre = fichier.add_paragraph()
    trace = titre.add_run(bloc.titre)
    trace.bold = True
    trace.font.size = Pt(10)

    maximum = bloc.maximum
    largeur_libelle = max(len(k) for k, _ in bloc.valeurs)
    for libelle, valeur in bloc.valeurs:
        pleins = int(round(valeur / maximum * LARGEUR_BARRE))
        ligne = fichier.add_paragraph()
        ligne.paragraph_format.space_after = Pt(0)
        graphie = ligne.add_run(
            f"{libelle.ljust(largeur_libelle)}  "
            f"{'█' * pleins}{'·' * (LARGEUR_BARRE - pleins)}  {_nombre(valeur)}"
        )
        graphie.font.name = "Consolas"
        graphie.font.size = Pt(8)
    fichier.add_paragraph()


def _encadre(fichier, bloc: Encadre) -> None:
    teinte = {"alerte": ROUGE_ALERTE, "attention": ORANGE_ATTENTION}.get(
        bloc.ton, VERT_ADMIN
    )
    cadre = fichier.add_table(rows=1, cols=1)
    cadre.style = "Table Grid"
    cellule = cadre.rows[0].cells[0]
    _fond(cellule, "FBFAF6")
    cellule.paragraphs[0].text = ""
    titre = cellule.paragraphs[0].add_run(bloc.titre)
    titre.bold = True
    titre.font.color.rgb = teinte
    corps = cellule.add_paragraph()
    corps.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    corps.add_run(bloc.texte).font.size = Pt(9.5)
    fichier.add_paragraph()


def _signature(fichier, document: Document) -> None:
    if document.mention_finale:
        mention = fichier.add_paragraph()
        mention.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        trace = mention.add_run(document.mention_finale)
        trace.italic = True
        trace.font.size = Pt(9)

    fichier.add_paragraph()
    for texte, gras in (
        (f"{document.lieu}, le {_date(document.etabli_le)}", False),
        (document.signataire, True),
    ):
        paragraphe = fichier.add_paragraph()
        paragraphe.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        trace = paragraphe.add_run(texte)
        trace.bold = gras


# ------------------------------------------------------------------ outils


def _fond(cellule, teinte: str) -> None:
    """python-docx n'expose pas la couleur de fond d'une cellule ; il faut
    poser l'élément XML correspondant à la main."""
    ombrage = OxmlElement("w:shd")
    ombrage.set(qn("w:val"), "clear")
    ombrage.set(qn("w:color"), "auto")
    ombrage.set(qn("w:fill"), teinte)
    cellule._tc.get_or_add_tcPr().append(ombrage)


def _filet(paragraphe) -> None:
    """Trait horizontal sous la titulature, posé en bordure de paragraphe."""
    bordures = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), "10")
    bas.set(qn("w:space"), "1")
    bas.set(qn("w:color"), "1B5E3A")
    bordures.append(bas)
    paragraphe._p.get_or_add_pPr().append(bordures)


def _nombre(valeur: float) -> str:
    return str(int(valeur)) if valeur == int(valeur) else f"{valeur:.1f}"


def _date(valeur: datetime | None) -> str:
    return valeur.strftime("%d/%m/%Y à %H h %M") if valeur else "—"
